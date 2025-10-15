# ================================================
# 🔹 SIMULADOR DE REGRESSÃO (Random Forest)
# ================================================

import pandas as pd
import numpy as np
import os
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches
from sklearn.preprocessing import StandardScaler
from sklearn.model_selection import cross_val_predict, KFold, cross_val_score
from sklearn.ensemble import RandomForestRegressor
from sklearn.metrics import r2_score, mean_squared_error, mean_absolute_error
from statsmodels.stats.outliers_influence import variance_inflation_factor

# ================================================
# 1. CONFIGURAÇÕES INICIAIS
# ================================================
csv_path = r"C:\Users\Ana Clara Azevedo\Documents\Univesp (2ª Graduação)\5º Semestre\Projeto Integrador em Computação II\Quinzena 5\Dados Target.csv"
pasta_relatorio = r"C:\Users\Ana Clara Azevedo\Documents\Univesp (2ª Graduação)\5º Semestre\Projeto Integrador em Computação II\Quinzena 5"
os.makedirs(pasta_relatorio, exist_ok=True)
relatorio_path = os.path.join(pasta_relatorio, "Análise Final.docx")

# ================================================
# 2. CARREGAR E PREPARAR DADOS
# ================================================
df = pd.read_csv(csv_path, sep=";", decimal=",")
for col in df.columns:
    if df[col].dtype == object:
        df[col] = df[col].str.replace(" ", "").astype(float)

# Removendo Inputs 5 e 6
df = df.drop(columns=["Input05", "Input06"])

y = df["Target"]
X_full = df.drop("Target", axis=1)

# ================================================
# 3. FUNÇÃO DE REMOÇÃO DE OUTLIERS POR PERCENTIL
# ================================================
def remover_outliers_percentil(df, lower=0.01, upper=0.99):
    df_out = df.copy()
    for col in df_out.columns:
        low_val = df_out[col].quantile(lower)
        high_val = df_out[col].quantile(upper)
        df_out = df_out[(df_out[col] >= low_val) & (df_out[col] <= high_val)]
    return df_out

df_tratado = remover_outliers_percentil(X_full.join(y))

y_tratado = df_tratado["Target"]
X_tratado = df_tratado.drop("Target", axis=1)

# ================================================
# 4. PADRONIZAÇÃO
# ================================================
scaler = StandardScaler()
X_scaled = pd.DataFrame(scaler.fit_transform(X_tratado), columns=X_tratado.columns)

# ================================================
# 5. RANDOM FOREST
# ================================================
cv = KFold(n_splits=10, shuffle=True, random_state=42)
rf_model = RandomForestRegressor(random_state=42, n_estimators=150, max_depth=12)
y_pred_cv = cross_val_predict(rf_model, X_scaled, y_tratado, cv=cv)

r2 = r2_score(y_tratado, y_pred_cv)
mae = mean_absolute_error(y_tratado, y_pred_cv)
rmse = np.sqrt(mean_squared_error(y_tratado, y_pred_cv))

# Treinar o modelo completo para gráficos de importância
rf_model.fit(X_scaled, y_tratado)
importancias = rf_model.feature_importances_

# ================================================
# 6. CRIAÇÃO DO RELATÓRIO
# ================================================
doc = Document()
doc.add_heading("Análise Preditiva - Random Forest", 0)
doc.add_paragraph(
    "Este relatório apresenta a análise preditiva utilizando Random Forest sobre a base "
    "com outliers removidos pelo Percentil 1%-99%."
)

# -----------------------------
# Predições x Valores Reais
# -----------------------------
plt.figure(figsize=(6,4))
plt.scatter(y_tratado, y_pred_cv, alpha=0.6)
plt.plot([y_tratado.min(), y_tratado.max()], [y_tratado.min(), y_tratado.max()], 'r--')
plt.xlabel("Valores reais")
plt.ylabel("Predições k-fold")
plt.title("Predições x Valores Reais - Random Forest")
plt.tight_layout()
path_pred = os.path.join(pasta_relatorio, "Predições x Valores Reais.png")
plt.savefig(path_pred, dpi=300)
plt.close()
doc.add_picture(path_pred, width=Inches(6))

# -----------------------------
# Distribuição de Resíduos
# -----------------------------
residuos = y_tratado - y_pred_cv
plt.figure(figsize=(6,4))
sns.histplot(residuos, kde=True)
plt.title("Distribuição de Resíduos - Random Forest")
plt.xlabel("Resíduo")
plt.ylabel("Frequência")
plt.tight_layout()
path_res = os.path.join(pasta_relatorio, "Distribuição de Resíduos.png")
plt.savefig(path_res, dpi=300)
plt.close()
doc.add_picture(path_res, width=Inches(6))

# -----------------------------
# Resíduos x Predições
# -----------------------------
plt.figure(figsize=(6,4))
plt.scatter(y_pred_cv, residuos, alpha=0.6)
plt.axhline(0, color='r', linestyle='--')
plt.xlabel("Predições k-fold")
plt.ylabel("Resíduos")
plt.title("Resíduos x Predições - Random Forest")
plt.tight_layout()
path_resid = os.path.join(pasta_relatorio, "Residuos x Prediçoes.png")
plt.savefig(path_resid, dpi=300)
plt.close()
doc.add_picture(path_resid, width=Inches(6))

# -----------------------------
# Curva de Aprendizado
# -----------------------------
n_trees = [10, 50, 100, 150, 200]
scores = []
for n in n_trees:
    rf = RandomForestRegressor(n_estimators=n, max_depth=12, random_state=42)
    r2_fold = cross_val_score(rf, X_scaled, y_tratado, cv=cv, scoring='r2')
    scores.append(r2_fold.mean())

plt.figure(figsize=(6,4))
plt.plot(n_trees, scores, marker='o')
plt.title("Curva de Aprendizado - Random Forest")
plt.xlabel("Número de árvores")
plt.ylabel("R² médio (10 folds)")
plt.tight_layout()
path_curve = os.path.join(pasta_relatorio, "Curva de Aprendizado - Random Forest.png")
plt.savefig(path_curve, dpi=300)
plt.close()
doc.add_picture(path_curve, width=Inches(6))

# -----------------------------
# Importância das Variáveis
# -----------------------------
indices = np.argsort(importancias)[::-1]
plt.figure(figsize=(8,6))
sns.barplot(x=importancias[indices], y=X_tratado.columns[indices])
plt.title("Importância das Variáveis - Random Forest")
plt.tight_layout()
path_import = os.path.join(pasta_relatorio, "Importância das Variáveis.png")
plt.savefig(path_import, dpi=300)
plt.close()
doc.add_picture(path_import, width=Inches(6))

# -----------------------------
# VIF x Importância
# -----------------------------
vif_vals = [variance_inflation_factor(X_scaled.values, i) for i in range(X_scaled.shape[1])]
features = X_tratado.columns

plt.figure(figsize=(8,5))
sns.barplot(x=importancias, y=features, color='b', label='Importância')
sns.scatterplot(x=vif_vals, y=features, color='r', s=100, label='VIF', marker='D')
plt.title("Importância x VIF - Random Forest")
plt.xlabel("Valor")
plt.legend()
plt.tight_layout()
path_fiv = os.path.join(pasta_relatorio, "VIF x Importância.png")
plt.savefig(path_fiv, dpi=300)
plt.close()
doc.add_picture(path_fiv, width=Inches(6))

# -----------------------------
# Métricas
# -----------------------------
doc.add_heading("Métricas de Desempenho - Random Forest", level=1)
table = doc.add_table(rows=2, cols=4)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Modelo"
hdr_cells[1].text = "R²"
hdr_cells[2].text = "MAE"
hdr_cells[3].text = "RMSE"

row = table.rows[1].cells
row[0].text = "Random Forest"
row[1].text = f"{r2:.4f}"
row[2].text = f"{mae:.2f}"
row[3].text = f"{rmse:.2f}"

# -----------------------------
# Conclusões
# -----------------------------
doc.add_heading("Conclusões Finais", level=1)
doc.add_paragraph(
    "- Random Forest apresenta bom desempenho preditivo interno.\n"
    "- Remoção de outliers pelo Percentil 1%-99% aumenta robustez.\n"
    "- Validação cruzada k-fold garante avaliação confiável.\n"
    "- Gráficos de importância das variáveis e VIF ajudam na interpretação do modelo.\n"
)

doc.save(relatorio_path)
print(f"✅ Relatório final salvo em: {relatorio_path}")
