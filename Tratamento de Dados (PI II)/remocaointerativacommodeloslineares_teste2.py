import pandas as pd
from sklearn.model_selection import train_test_split
from sklearn.linear_model import LinearRegression, Ridge, Lasso, ElasticNet
from sklearn.decomposition import PCA
from docx import Document

# -----------------------------
# 1. Leitura do CSV
# -----------------------------
csv_path = r"C:\Users\BRPiresAn3\Downloads\Tratamento de Dados - PI II\Dados Target.csv"
df = pd.read_csv(csv_path, sep=";", decimal=",")  # vírgula como decimal
print("Dados carregados:")
print(df.head())

# Corrigir possíveis espaços nos números
for col in df.columns:
    if df[col].dtype == object:
        df[col] = df[col].str.replace(" ", "").astype(float)

# Separar X e y
X = df.drop(columns=["Target"])
y = df["Target"]

# -----------------------------
# 2. Função para treinar modelos
# -----------------------------
def treinar_modelos(X, y):
    resultados = {}
    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)
    
    # Linear
    lr = LinearRegression().fit(X_train, y_train)
    resultados["LinearRegression"] = lr.score(X_test, y_test)
    
    # Ridge
    ridge = Ridge().fit(X_train, y_train)
    resultados["Ridge"] = ridge.score(X_test, y_test)
    
    # Lasso
    lasso = Lasso(max_iter=10000).fit(X_train, y_train)
    resultados["Lasso"] = lasso.score(X_test, y_test)
    
    # ElasticNet
    enet = ElasticNet(max_iter=10000).fit(X_train, y_train)
    resultados["ElasticNet"] = enet.score(X_test, y_test)
    
    # PCA + LinearRegression
    pca = PCA(n_components=min(X.shape[1], X_train.shape[0])).fit(X_train)
    X_train_pca = pca.transform(X_train)
    X_test_pca = pca.transform(X_test)
    lr_pca = LinearRegression().fit(X_train_pca, y_train)
    resultados["PCA_LinearRegression"] = lr_pca.score(X_test_pca, y_test)
    
    return resultados

# -----------------------------
# 3. Resultados com todos os inputs
# -----------------------------
resultados_originais = treinar_modelos(X, y)
print("R² com todos os inputs:")
print(resultados_originais)

# -----------------------------
# 4. Testar remoção de cada input
# -----------------------------
remocoes = {}
for col in X.columns:
    X_temp = X.drop(columns=[col])
    resultados_temp = treinar_modelos(X_temp, y)
    remocoes[col] = resultados_temp

print("Resultados removendo cada input:")
for col, res in remocoes.items():
    print(f"Removendo {col}: {res}")

# -----------------------------
# 5. Gerar relatório Word
# -----------------------------
doc = Document()
doc.add_heading("Comparativo entre Modelos", level=0)

doc.add_heading("1. Resultados com todos os inputs", level=1)
table = doc.add_table(rows=1, cols=2)
table.style = "Light List"
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Modelo"
hdr_cells[1].text = "R²"
for modelo, r2 in resultados_originais.items():
    row_cells = table.add_row().cells
    row_cells[0].text = modelo
    row_cells[1].text = f"{r2:.4f}"

doc.add_heading("2. Resultados removendo cada input", level=1)
for col, res in remocoes.items():
    doc.add_paragraph(f"Removendo {col}:")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light List"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Modelo"
    hdr_cells[1].text = "R²"
    for modelo, r2 in res.items():
        row_cells = table.add_row().cells
        row_cells[0].text = modelo
        row_cells[1].text = f"{r2:.4f}"
    doc.add_paragraph("")  # espaço entre tabelas

# Justificativa
doc.add_heading("3. Justificativa das mudanças no R²", level=1)
doc.add_paragraph(
    "A remoção de variáveis altamente colineares pode reduzir o risco de instabilidade nos coeficientes "
    "e melhora a interpretabilidade dos modelos. No entanto, a exclusão de variáveis também pode causar "
    "queda no R², pois informações úteis para predição podem ser eliminadas. "
    "Como os inputs são confidenciais, não sabemos quais são mais relevantes, então a avaliação é feita "
    "principalmente pelo desempenho (R²) dos modelos."
)

# Salvar relatório
relatorio_path = r"C:\Users\BRPiresAn3\Downloads\Tratamento de Dados - PI II\Comparativo com Modelos Lineares.docx"
doc.save(relatorio_path)
print(f"Relatório salvo em: {relatorio_path}")

