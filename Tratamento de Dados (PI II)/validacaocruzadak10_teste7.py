import pandas as pd
import numpy as np
import os
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from sklearn.preprocessing import StandardScaler
from sklearn.model_selection import cross_val_score, cross_val_predict, KFold
from sklearn.linear_model import LinearRegression, Ridge, Lasso, ElasticNet
from sklearn.ensemble import RandomForestRegressor, GradientBoostingRegressor
from sklearn.metrics import r2_score, mean_squared_error, mean_absolute_error
import seaborn as sns

# --- 1. Carregar dados ---
csv_path = r"C:\Users\Ana Clara Azevedo\Documents\Univesp (2ª Graduação)\5º Semestre\Projeto Integrador em Computação II\Quinzena 5\Dados Target.csv"
df = pd.read_csv(csv_path, sep=";", decimal=",")

for col in df.columns:
    if df[col].dtype == object:
        df[col] = df[col].str.replace(" ", "").astype(float)

y = df["Target"]
X = df.drop("Target", axis=1)

# --- 2. Padronização ---
scaler = StandardScaler()
X_scaled = pd.DataFrame(scaler.fit_transform(X), columns=X.columns)

# --- 3. Modelos ---
modelos = {
    "LinearRegression": LinearRegression(),
    "Ridge": Ridge(),
    "Lasso": Lasso(max_iter=10000),
    "ElasticNet": ElasticNet(max_iter=10000),
    "RandomForest": RandomForestRegressor(random_state=42),
    "GradientBoosting": GradientBoostingRegressor(random_state=42)
}

# --- 4. Validação cruzada ---
cv = KFold(n_splits=10, shuffle=True, random_state=42)
resultados_cv = {}

for nome, modelo in modelos.items():
    y_pred_cv = cross_val_predict(modelo, X_scaled, y, cv=cv)
    r2 = r2_score(y, y_pred_cv)
    mae = mean_absolute_error(y, y_pred_cv)
    rmse = np.sqrt(mean_squared_error(y, y_pred_cv))
    resultados_cv[nome] = {"R2": r2, "MAE": mae, "RMSE": rmse}

# --- 5. Criar pasta e iniciar documento Word ---
pasta_relatorio = r"C:\Users\BRPiresAn3\Downloads\Tratamento de Dados - PI II"
os.makedirs(pasta_relatorio, exist_ok=True)
relatorio_path = os.path.join(pasta_relatorio, "Validação Cruzada - Teste 7.docx")
doc = Document()
doc.add_heading("Análise Completa de Modelos Preditivos", 0)

# --- 6. Introdução ---
doc.add_heading("1. Introdução", level=1)
doc.add_paragraph(
    "Esta análise expande a versão anterior utilizando Validação Cruzada, múltiplas métricas de erro "
    "(R², MAE, RMSE), análise de importância das variáveis e gráficos de resíduos. "
    "Isso fornece uma visão mais completa e robusta do desempenho dos modelos."
    "A validação cruzada geralmente dá uma avaliação mais confiável e robusta que o simples holdout 70/30 porque:"
    "- No holdout, você depende de uma única divisão aleatória dos dados. Se essa divisão for “azarada” (exemplo: poucos casos de um tipo, distribuição diferente), a avaliação pode ficar enviesada."
    "- Na validação cruzada, o modelo é testado em várias “rodadas” diferentes, cada uma usando partes diferentes dos dados para teste e treino. Isso suaviza flutuações e dá uma média mais estável da performance."

)

# --- 7. Tabela de Resultados ---
doc.add_heading("2. Avaliação com Validação Cruzada", level=1)
table = doc.add_table(rows=1, cols=4)
table.style = "Light List"
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Modelo"
hdr_cells[1].text = "R²"
hdr_cells[2].text = "MAE"
hdr_cells[3].text = "RMSE"

for modelo, res in resultados_cv.items():
    row_cells = table.add_row().cells
    row_cells[0].text = modelo
    row_cells[1].text = f"{res['R2']:.4f}"
    row_cells[2].text = f"{res['MAE']:.2f}"
    row_cells[3].text = f"{res['RMSE']:.2f}"

# --- 8. Gráfico de Métricas ---
doc.add_heading("3. Comparativo Visual das Métricas", level=1)
df_resultados = pd.DataFrame(resultados_cv).T
metrica_plot_path = os.path.join(pasta_relatorio, "Gráfico das Métricas - Teste 7.png")
df_resultados.plot(kind="barh", figsize=(10, 6))
plt.title("Comparação das Métricas dos Modelos")
plt.tight_layout()
plt.savefig(metrica_plot_path, dpi=300)
plt.close()
doc.add_picture(metrica_plot_path, width=Inches(6))

# --- 9. Importância das Variáveis (Random Forest) ---
rf_model = RandomForestRegressor(random_state=42)
rf_model.fit(X_scaled, y)
importancias = rf_model.feature_importances_
indices = np.argsort(importancias)[::-1]
plt.figure(figsize=(8,6))
sns.barplot(x=importancias[indices], y=X.columns[indices])
plt.title("Importância das Variáveis - Random Forest")
plt.tight_layout()
importancia_path = os.path.join(pasta_relatorio, "Importância das Variáveis - Teste 7.png")
plt.savefig(importancia_path, dpi=300)
plt.close()

doc.add_heading("4. Importância das Variáveis", level=1)
doc.add_paragraph(
    "Abaixo está o gráfico com as variáveis mais relevantes segundo o modelo Random Forest. "
    "Isso ajuda a entender quais variáveis mais contribuem para explicar o target."
)
doc.add_picture(importancia_path, width=Inches(6))

# --- 10. Análise de Resíduos (Gradient Boosting) ---
gb_model = GradientBoostingRegressor(random_state=42)
y_pred_gb = cross_val_predict(gb_model, X_scaled, y, cv=cv)
residuos = y - y_pred_gb
plt.figure(figsize=(8,6))
sns.scatterplot(x=y_pred_gb, y=residuos)
plt.axhline(0, color='red', linestyle='--')
plt.xlabel("Valores Previstos")
plt.ylabel("Resíduos")
plt.title("Resíduos x Previsões - Gradient Boosting")
plt.tight_layout()
residuos_path = os.path.join(pasta_relatorio, "Gráfico de Resíduos - Teste 7.png")
plt.savefig(residuos_path, dpi=300)
plt.close()

doc.add_heading("5. Análise de Resíduos", level=1)
doc.add_paragraph(
    "O gráfico de resíduos ajuda a identificar padrões nos erros. "
    "Idealmente, os resíduos devem estar distribuídos aleatoriamente em torno de zero, "
    "o que indica que o modelo não está cometendo erros sistemáticos."
)
doc.add_picture(residuos_path, width=Inches(6))

# --- 11. Conclusão ---
doc.add_heading("6. Conclusões", level=1)
doc.add_paragraph(
    "- A validação cruzada fornece uma estimativa mais confiável do desempenho dos modelos.\n"
    "- Métricas como MAE e RMSE complementam o R² ao mostrar o erro médio do modelo.\n"
    "- A importância das variáveis ajuda a interpretar o modelo.\n"
    "- A análise de resíduos é essencial para verificar se o modelo está cometendo erros sistemáticos.\n"
    "- Modelos ensemble (Random Forest, Gradient Boosting) continuam sendo os mais eficazes.\n"
)

# --- 12. Salvar relatório ---
doc.save(relatorio_path)
print(f"Relatório completo salvo em: {relatorio_path}")
