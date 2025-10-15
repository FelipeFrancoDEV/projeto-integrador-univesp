import pandas as pd
import numpy as np
from sklearn.preprocessing import StandardScaler
from sklearn.model_selection import train_test_split
from sklearn.linear_model import LinearRegression, Ridge, Lasso, ElasticNet
from sklearn.ensemble import RandomForestRegressor, GradientBoostingRegressor
from sklearn.metrics import r2_score
from statsmodels.stats.outliers_influence import variance_inflation_factor
from docx import Document
import os

# --- Caminho dos dados ---
csv_path = r"C:\Users\BRPiresAn3\Downloads\Tratamento de Dados - PI II\Dados Target.csv"
df = pd.read_csv(csv_path, sep=";", decimal=",")

# --- Separar target e inputs ---
y = df["Target"]
X = df.drop("Target", axis=1)

# --- Padronização ---
scaler = StandardScaler()
X_scaled = pd.DataFrame(scaler.fit_transform(X), columns=X.columns)

# --- Função para calcular VIF ---
def calcular_vif(X):
    vif_data = pd.DataFrame()
    vif_data["feature"] = X.columns
    vif_data["VIF"] = [variance_inflation_factor(X.values, i) for i in range(X.shape[1])]
    return vif_data

# --- Remover variáveis com VIF alto ---
vif_limite = 1000
X_vif = X_scaled.copy()
vif = calcular_vif(X_vif)
variaveis_removidas = []

while vif["VIF"].max() > vif_limite:
    max_vif_feature = vif.sort_values("VIF", ascending=False).iloc[0]["feature"]
    X_vif = X_vif.drop(columns=[max_vif_feature])
    variaveis_removidas.append(max_vif_feature)
    vif = calcular_vif(X_vif)

# --- Treinar modelos ---
def treinar_modelos(X, y):
    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.3, random_state=42)
    
    modelos = {
        "LinearRegression": LinearRegression(),
        "Ridge": Ridge(),
        "Lasso": Lasso(),
        "ElasticNet": ElasticNet(),
        "RandomForest": RandomForestRegressor(random_state=42),
        "GradientBoosting": GradientBoostingRegressor(random_state=42)
    }
    
    resultados = {}
    
    for nome, modelo in modelos.items():
        modelo.fit(X_train, y_train)
        y_pred = modelo.predict(X_test)
        resultados[nome] = r2_score(y_test, y_pred)
        
    return resultados

# --- Seleção de features pelo Lasso ---
lasso = Lasso(alpha=0.01)
lasso.fit(X_vif, y)
features_selecionadas = X_vif.columns[lasso.coef_ != 0].tolist()
X_final = X_vif[features_selecionadas]

# --- Treinar modelos com features finais ---
resultados = treinar_modelos(X_final, y)

# --- Criar relatório Word detalhado ---
pasta_relatorio = r"C:\Users\BRPiresAn3\Downloads\Tratamento de Dados - PI II"
os.makedirs(pasta_relatorio, exist_ok=True)
relatorio_path = os.path.join(pasta_relatorio, "Remoção do Input 5.docx")

doc = Document()
doc.add_heading("Comparativo de Modelos - Tratamento de Colinearidade", 0)

# Variáveis removidas
doc.add_heading("Variáveis removidas por alta correlação (VIF):", level=1)
doc.add_paragraph(
    "Variáveis com VIF maior que {} foram removidas para reduzir multicolinearidade e "
    "evitar instabilidade nos coeficientes. Isto melhora a interpretabilidade, "
    "mas pode reduzir ligeiramente o desempenho (R²).".format(vif_limite)
)
doc.add_paragraph(str(variaveis_removidas))

# Features finais pelo Lasso
doc.add_heading("Features finais selecionadas pelo Lasso:", level=1)
doc.add_paragraph(
    "O Lasso realiza regularização L1, forçando coeficientes irrelevantes a zero, "
    "selecionando apenas as features mais importantes para predição."
)
doc.add_paragraph(str(features_selecionadas))

# R² dos modelos
doc.add_heading("R² médio dos modelos:", level=1)
doc.add_paragraph(
    "R² indica a proporção da variância do target explicada pelo modelo. "
    "Modelos lineares podem ser afetados por multicolinearidade, "
    "enquanto modelos de ensemble como Random Forest e Gradient Boosting "
    "lidam melhor com interações complexas entre features."
)
for modelo, r2 in resultados.items():
    doc.add_paragraph(f"{modelo}: {r2:.4f}")

# Explicações das técnicas usadas
doc.add_heading("Explicação das técnicas aplicadas:", level=1)
doc.add_paragraph(
    "1. Padronização (StandardScaler): coloca todas as features na mesma escala, "
    "evitando que variáveis com magnitudes maiores dominem o modelo.\n\n"
    "2. VIF (Variance Inflation Factor): identifica colinearidade entre variáveis, "
    "removendo aquelas que trazem redundância excessiva.\n\n"
    "3. Lasso: regularização L1 que ajuda na seleção automática de features, "
    "reduzindo overfitting.\n\n"
    "4. Modelos treinados: LinearRegression, Ridge, Lasso, ElasticNet, RandomForest, "
    "GradientBoosting para comparar desempenho e robustez.\n\n"
    "5. Avaliação: R² é usado como métrica principal, já que não conhecemos o significado "
    "de cada input e não podemos avaliar relevância individual das features."
)

doc.save(relatorio_path)
print(f"Relatório detalhado salvo em: {relatorio_path}")
