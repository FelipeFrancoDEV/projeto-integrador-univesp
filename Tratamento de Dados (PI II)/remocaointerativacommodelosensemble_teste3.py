import pandas as pd
import numpy as np
from sklearn.preprocessing import StandardScaler
from sklearn.model_selection import train_test_split
from sklearn.linear_model import LinearRegression, Ridge, Lasso, ElasticNet
from sklearn.ensemble import RandomForestRegressor, GradientBoostingRegressor
from sklearn.metrics import r2_score
from docx import Document
import os

# --- Caminho dos dados ---
csv_path = r"C:\Users\BRPiresAn3\Downloads\Tratamento de Dados - PI II\Dados Target.csv"
df = pd.read_csv(csv_path, sep=";", decimal=",")  # vírgula como decimal
print("Dados carregados:")
print(df.head())

# Corrigir possíveis espaços nos números
for col in df.columns:
    if df[col].dtype == object:
        df[col] = df[col].str.replace(" ", "").astype(float)

# Separar target e inputs
y = df["Target"]
X = df.drop("Target", axis=1)

# Padronização
scaler = StandardScaler()
X_scaled = pd.DataFrame(scaler.fit_transform(X), columns=X.columns)

# Função para treinar modelos
def treinar_modelos(X, y):
    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.3, random_state=42)
    
    modelos = {
        "LinearRegression": LinearRegression(),
        "Ridge": Ridge(),
        "Lasso": Lasso(max_iter=10000),
        "ElasticNet": ElasticNet(max_iter=10000),
        "RandomForest": RandomForestRegressor(random_state=42),
        "GradientBoosting": GradientBoostingRegressor(random_state=42)
    }
    
    resultados = {}
    for nome, modelo in modelos.items():
        modelo.fit(X_train, y_train)
        y_pred = modelo.predict(X_test)
        resultados[nome] = r2_score(y_test, y_pred)
    return resultados

# Resultados com todos os inputs
resultados_originais = treinar_modelos(X_scaled, y)

# Testar remoção de cada input individualmente
remocoes = {}
for col in X_scaled.columns:
    X_temp = X_scaled.drop(columns=[col])
    resultados_temp = treinar_modelos(X_temp, y)
    remocoes[col] = resultados_temp

# Criar relatório Word
pasta_relatorio = r"C:\Users\BRPiresAn3\Downloads\Tratamento de Dados - PI II"
os.makedirs(pasta_relatorio, exist_ok=True)
relatorio_path = os.path.join(pasta_relatorio, "Comparativo com Modelos Ensemble.docx")

doc = Document()
doc.add_heading("Comparativo de Modelos - Análise de Impacto de Variáveis", 0)

# 1. Resultados com todos os inputs
doc.add_heading("1. Resultados com todos os inputs", level=1)
table = doc.add_table(rows=1, cols=2)
table.style = "Light List"
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Modelo"
hdr_cells[1].text = "R²"
for modelo, r2 in resultados_originais.items():
    row_cells = table.add_row().cells
    row_cells[0].text = modelo
    row_cells[1].text = f"{r2:.4f}"

# 2. Resultados removendo cada input
doc.add_heading("2. Resultados removendo cada input", level=1)
for col, res in remocoes.items():
    doc.add_paragraph(f"Removendo {col}:")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light List"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Modelo"
    hdr_cells[1].text = "R²"
    for modelo, r2 in res.items():
        row_cells = table.add_row().cells
        row_cells[0].text = modelo
        row_cells[1].text = f"{r2:.4f}"
    doc.add_paragraph("")  # espaço entre tabelas

# 3. Justificativa e explicações
doc.add_heading("3. Interpretação dos Resultados", level=1)
doc.add_paragraph(
    "A análise teve como objetivo comparar o desempenho de diferentes modelos de regressão "
    "ao prever o Target, considerando a possível colinearidade entre os inputs.\n\n"
    "Explicações detalhadas:\n"
    "1. Padronização (StandardScaler): coloca todas as features na mesma escala para que nenhuma domine o modelo.\n"
    "2. Treinamento de modelos lineares: LinearRegression, Ridge, Lasso, ElasticNet.\n"
    "3. Treinamento de modelos de ensemble: RandomForest, GradientBoosting, que lidam melhor com interações complexas.\n"
    "4. Avaliação pelo R²: mostra a proporção da variação do Target explicada pelo modelo.\n"
    "5. Teste de remoção de cada input: avalia o efeito isolado de cada variável no desempenho. Pequenas alterações no R² indicam que a variável tem impacto baixo isoladamente, mas o efeito conjunto ainda pode ser relevante.\n"
    "6. Considerações sobre variáveis confidenciais: não sabemos o significado de cada input, então a avaliação é feita principalmente pelo desempenho dos modelos, garantindo uma análise imparcial.\n\n"
    "Observações:\n"
    "- Se o R² melhora levemente ao remover um input, isso pode indicar colinearidade ou informação redundante.\n"
    "- Se o R² cai drasticamente, o input é importante para a predição.\n"
    "- Modelos de ensemble geralmente apresentam maior robustez frente a multicolinearidade."
)

doc.save(relatorio_path)
print(f"Relatório detalhado salvo em: {relatorio_path}")
