import pandas as pd
import numpy as np
import os
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from sklearn.preprocessing import StandardScaler
from sklearn.model_selection import train_test_split
from sklearn.linear_model import LinearRegression, Ridge, Lasso, ElasticNet
from sklearn.ensemble import RandomForestRegressor, GradientBoostingRegressor
from sklearn.metrics import r2_score

# --- Caminho dos dados ---
csv_path = r"C:\Users\BRPiresAn3\Downloads\Tratamento de Dados - PI II\Dados Target.csv"
df = pd.read_csv(csv_path, sep=";", decimal=",")

# Corrigir possíveis espaços nos números
for col in df.columns:
    if df[col].dtype == object:
        df[col] = df[col].str.replace(" ", "").astype(float)

# --- Separar target e inputs ---
y = df["Target"]
X = df.drop("Target", axis=1)

# --- Padronização ---
scaler = StandardScaler()
X_scaled = pd.DataFrame(scaler.fit_transform(X), columns=X.columns)

# --- Divisão treino/teste 80/20 ---
X_train, X_test, y_train, y_test = train_test_split(X_scaled, y, test_size=0.2, random_state=42)

# --- Treinar modelos ---
modelos = {
    "LinearRegression": LinearRegression(),
    "Ridge": Ridge(),
    "Lasso": Lasso(max_iter=10000),
    "ElasticNet": ElasticNet(max_iter=10000),
    "RandomForest": RandomForestRegressor(random_state=42),
    "GradientBoosting": GradientBoostingRegressor(random_state=42)
}

resultados = {}
for nome, modelo in modelos.items():
    modelo.fit(X_train, y_train)
    y_pred = modelo.predict(X_test)
    resultados[nome] = r2_score(y_test, y_pred)

# --- Criar relatório Word ---
pasta_relatorio = r"C:\Users\BRPiresAn3\Downloads\Tratamento de Dados - PI II"
os.makedirs(pasta_relatorio, exist_ok=True)
relatorio_path = os.path.join(pasta_relatorio, "Pareto.docx")

doc = Document()
doc.add_heading("Comparativo de Modelos - Avaliação 80/20", 0)

# Resultados com todos os inputs
doc.add_heading("1. Resultados com todos os inputs", level=1)
doc.add_paragraph(
    "Todos os inputs foram utilizados juntos. A divisão 80/20 garante que o modelo seja avaliado "
    "em dados que ele não viu durante o treino, simulando como ele se comportaria na prática."
)
table = doc.add_table(rows=1, cols=2)
table.style = "Light List"
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Modelo"
hdr_cells[1].text = "R²"
for modelo, r2 in resultados.items():
    row_cells = table.add_row().cells
    row_cells[0].text = modelo
    row_cells[1].text = f"{r2:.4f}"

# Explicação das técnicas aplicadas
doc.add_heading("2. Explicação das técnicas aplicadas", level=1)
doc.add_paragraph(
    "1. Padronização (StandardScaler): coloca todas as features na mesma escala, "
    "evitando que variáveis com magnitudes maiores dominem a análise.\n\n"
    "2. Treino/Teste (80/20): 80% dos dados para treino, 20% para teste, garantindo avaliação justa.\n\n"
    "3. Modelos Lineares: LinearRegression, Ridge, Lasso, ElasticNet, funcionam melhor quando há relações aproximadamente lineares entre as variáveis.\n\n"
    "4. Modelos de Ensemble: RandomForest, GradientBoosting. Combinam várias 'mini decisões' (árvores de decisão) para melhorar a previsão. "
    "Podem capturar relações lineares e não lineares. Analogia simples: é como consultar vários especialistas e tirar uma média das opiniões.\n\n"
    "5. R²: indica a proporção da variância do target explicada pelo modelo. Quanto mais próximo de 1, melhor o modelo consegue explicar o comportamento dos dados."
)

# Interpretação dos resultados
doc.add_heading("3. Interpretação dos resultados", level=1)
doc.add_paragraph(
    "- Modelos lineares: R² em torno de 0,57 indica que cerca de 57% da variância do target é explicada pelas variáveis lineares.\n"
    "- Modelos ensemble: R² acima de 0,91 indica uma previsão muito mais precisa, pois eles lidam com relações complexas entre variáveis que os modelos lineares não capturam.\n"
    "- Avaliação conjunta (80/20) dá uma visão mais realista do desempenho do modelo no conjunto completo."
)

# --- Gráfico comparativo ---
cores = ['skyblue', 'skyblue', 'skyblue', 'skyblue', 'lightgreen', 'lightgreen']
plt.figure(figsize=(8,5))
plt.barh(list(resultados.keys()), list(resultados.values()), color=cores)
plt.xlabel("R²")
plt.title("Comparativo de R² entre modelos")
plt.xlim(0,1)
for i, v in enumerate(resultados.values()):
    plt.text(v + 0.01, i, f"{v:.2f}")
plt.tight_layout()
grafico_path = os.path.join(pasta_relatorio, "grafico_r2.png")
plt.savefig(grafico_path, dpi=300)
plt.close()

# Inserir gráfico no Word
doc.add_heading("4. Visualização do R²", level=1)
doc.add_picture(grafico_path, width=Inches(6))
doc.add_paragraph(
    "O gráfico mostra visualmente a performance de cada modelo. "
    "Modelos ensemble (RandomForest e GradientBoosting) apresentam R² mais alto, "
    "indicando que conseguem capturar relações mais complexas entre as variáveis, "
    "enquanto os modelos lineares explicam apenas a parte linear do comportamento do target."
)

# --- Salvar relatório ---
doc.save(relatorio_path)
print(f"Relatório completo salvo em: {relatorio_path}")
